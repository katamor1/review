from __future__ import annotations

from pathlib import Path
import zipfile

import pytest
from docx import Document

from bzr_step_count import word_document_service
from bzr_step_count.word_document_service import count_word_document


APP_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
    xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
  {pages}
</Properties>
"""


def _write_docx(path: Path, text: str = "ABC あい う") -> None:
    document = Document()
    document.add_paragraph(text)
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "表 1"
    document.save(path)


def _replace_app_xml(path: Path, pages_xml: str | None) -> None:
    replacement = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(replacement, "w") as target:
        for item in source.infolist():
            if item.filename == "docProps/app.xml":
                continue
            target.writestr(item, source.read(item.filename))
        if pages_xml is not None:
            target.writestr("docProps/app.xml", APP_XML.format(pages=pages_xml))
    replacement.replace(path)


def test_count_word_document_estimates_pages_and_keeps_metadata_as_reference(tmp_path, monkeypatch):
    docx_path = tmp_path / "sample.docx"
    _write_docx(docx_path, "A" * 7493)
    _replace_app_xml(docx_path, "<Pages>1</Pages>")
    monkeypatch.setattr(word_document_service, "_render_docx_to_pdf_page_count", lambda path: (None, ["no soffice"]))

    stats = count_word_document(docx_path)

    assert stats.path == str(docx_path)
    assert stats.display_page_count == 6
    assert stats.page_count_source == "estimated"
    assert stats.metadata_page_count == 1
    assert stats.estimated_page_count == 6
    assert stats.character_count_without_whitespace == 7495
    assert stats.character_count_with_whitespace == 7497
    assert "DOCXメタデータページ数 1 と採用ページ数 6 が一致しません" in stats.warnings
    assert "実ページ数を再計算できないため推定ページ数を採用しました" in stats.warnings


def test_count_word_document_warns_when_pages_metadata_is_missing(tmp_path):
    docx_path = tmp_path / "missing-pages.docx"
    _write_docx(docx_path)
    _replace_app_xml(docx_path, "")

    stats = count_word_document(docx_path)

    assert stats.metadata_page_count is None
    assert stats.page_count_source == "estimated"
    assert "ページ数メタデータが見つかりません: docProps/app.xml Pages" in stats.warnings


def test_count_word_document_warns_when_pages_metadata_is_invalid(tmp_path):
    docx_path = tmp_path / "invalid-pages.docx"
    _write_docx(docx_path)
    _replace_app_xml(docx_path, "<Pages>not-a-number</Pages>")

    stats = count_word_document(docx_path)

    assert stats.metadata_page_count is None
    assert stats.page_count_source == "estimated"
    assert "ページ数メタデータが数値ではありません: not-a-number" in stats.warnings


def test_count_word_document_prefers_rendered_pdf_page_count(tmp_path, monkeypatch):
    docx_path = tmp_path / "rendered.docx"
    _write_docx(docx_path, "A" * 5000)
    _replace_app_xml(docx_path, "<Pages>1</Pages>")
    monkeypatch.setattr(word_document_service, "_render_docx_to_pdf_page_count", lambda path: (4, []))

    stats = count_word_document(docx_path)

    assert stats.display_page_count == 4
    assert stats.page_count_source == "rendered_pdf"
    assert stats.estimated_page_count == 4
    assert stats.metadata_page_count == 1


def test_count_word_document_rejects_invalid_chars_per_page(tmp_path):
    docx_path = tmp_path / "sample.docx"
    _write_docx(docx_path)

    with pytest.raises(ValueError, match="1ページあたり文字数"):
        count_word_document(docx_path, chars_per_page=0)


def test_count_word_document_rejects_unsupported_or_unreadable_files(tmp_path):
    legacy_doc = tmp_path / "legacy.doc"
    legacy_doc.write_bytes(b"not a docx")
    missing_docx = tmp_path / "missing.docx"
    broken_docx = tmp_path / "broken.docx"
    broken_docx.write_bytes(b"not a zip")

    with pytest.raises(ValueError, match=r"\.docx のみ対応"):
        count_word_document(legacy_doc)
    with pytest.raises(ValueError, match="Word文書が見つかりません"):
        count_word_document(missing_docx)
    with pytest.raises(ValueError, match="Word文書を読み取れません"):
        count_word_document(broken_docx)
