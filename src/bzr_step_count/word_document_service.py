from __future__ import annotations

from dataclasses import dataclass, field
import math
from pathlib import Path
import shutil
import subprocess
import tempfile
import xml.etree.ElementTree as ET
import zipfile

from docx import Document
from pypdf import PdfReader


APP_PROPERTIES = "docProps/app.xml"
EXTENDED_PROPERTIES_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}"
DEFAULT_CHARS_PER_PAGE = 1400


@dataclass
class WordDocumentStats:
    path: str
    display_page_count: int | None
    page_count_source: str
    metadata_page_count: int | None
    estimated_page_count: int | None
    character_count_without_whitespace: int
    character_count_with_whitespace: int
    chars_per_page: int
    warnings: list[str] = field(default_factory=list)


def count_word_document(path: str | Path, *, chars_per_page: int = DEFAULT_CHARS_PER_PAGE) -> WordDocumentStats:
    if chars_per_page <= 0:
        raise ValueError("1ページあたり文字数は1以上を指定してください")

    document_path = Path(path)
    if document_path.suffix.lower() != ".docx":
        raise ValueError(f"Word文書は .docx のみ対応しています: {document_path}")
    if not document_path.exists() or not document_path.is_file():
        raise ValueError(f"Word文書が見つかりません: {document_path}")

    try:
        document = Document(document_path)
        text = "\n".join(_collect_text_blocks(document))
        metadata_page_count, metadata_warnings = _read_metadata_page_count(document_path)
    except Exception as exc:
        raise ValueError(f"Word文書を読み取れません: {document_path}") from exc

    character_count_without_whitespace = sum(1 for character in text if not character.isspace())
    character_count_with_whitespace = len(text)
    estimated_page_count = _estimate_page_count(character_count_without_whitespace, chars_per_page)
    rendered_page_count, render_warnings = _render_docx_to_pdf_page_count(document_path)
    warnings = [*metadata_warnings, *render_warnings]

    if rendered_page_count is not None:
        display_page_count = rendered_page_count
        page_count_source = "rendered_pdf"
    elif estimated_page_count is not None:
        display_page_count = estimated_page_count
        page_count_source = "estimated"
        warnings.append("実ページ数を再計算できないため推定ページ数を採用しました")
    else:
        display_page_count = None
        page_count_source = "unavailable"
        warnings.append("文字数が0のため推定ページ数を算出できません")

    if (
        metadata_page_count is not None
        and display_page_count is not None
        and metadata_page_count != display_page_count
    ):
        warnings.append(f"DOCXメタデータページ数 {metadata_page_count} と採用ページ数 {display_page_count} が一致しません")

    return WordDocumentStats(
        path=str(document_path),
        display_page_count=display_page_count,
        page_count_source=page_count_source,
        metadata_page_count=metadata_page_count,
        estimated_page_count=estimated_page_count,
        character_count_without_whitespace=character_count_without_whitespace,
        character_count_with_whitespace=character_count_with_whitespace,
        chars_per_page=chars_per_page,
        warnings=warnings,
    )


def _collect_text_blocks(document: Document) -> list[str]:
    blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        blocks.extend(_collect_table_text_blocks(table))
    return blocks


def _collect_table_text_blocks(table: object) -> list[str]:
    blocks: list[str] = []
    for row in table.rows:
        for cell in row.cells:
            blocks.extend(paragraph.text for paragraph in cell.paragraphs if paragraph.text)
            for nested_table in cell.tables:
                blocks.extend(_collect_table_text_blocks(nested_table))
    return blocks


def _estimate_page_count(character_count_without_whitespace: int, chars_per_page: int) -> int | None:
    if character_count_without_whitespace <= 0:
        return None
    return math.ceil(character_count_without_whitespace / chars_per_page)


def _read_metadata_page_count(path: Path) -> tuple[int | None, list[str]]:
    try:
        with zipfile.ZipFile(path, "r") as archive:
            try:
                app_xml = archive.read(APP_PROPERTIES)
            except KeyError:
                return None, [f"ページ数メタデータが見つかりません: {APP_PROPERTIES} Pages"]
    except zipfile.BadZipFile as exc:
        raise ValueError(f"Word文書を読み取れません: {path}") from exc

    try:
        root = ET.fromstring(app_xml)
    except ET.ParseError:
        return None, [f"ページ数メタデータを解析できません: {APP_PROPERTIES}"]

    pages = root.find(f"{EXTENDED_PROPERTIES_NS}Pages")
    if pages is None or pages.text is None or not pages.text.strip():
        return None, [f"ページ数メタデータが見つかりません: {APP_PROPERTIES} Pages"]

    raw_page_count = pages.text.strip()
    try:
        page_count = int(raw_page_count)
    except ValueError:
        return None, [f"ページ数メタデータが数値ではありません: {raw_page_count}"]
    return page_count, []


def _render_docx_to_pdf_page_count(path: Path) -> tuple[int | None, list[str]]:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None, ["LibreOffice soffice が見つからないためPDF変換によるページ数再計算をスキップしました"]

    with tempfile.TemporaryDirectory(prefix="review-stats-docx-") as temp_dir:
        output_dir = Path(temp_dir)
        args = [
            soffice,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_dir),
            str(path),
        ]
        try:
            completed = subprocess.run(
                args,
                shell=False,
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                timeout=120,
            )
        except (OSError, subprocess.TimeoutExpired) as exc:
            return None, [f"PDF変換によるページ数再計算に失敗しました: {exc}"]
        if completed.returncode != 0:
            message = completed.stderr.strip() or completed.stdout.strip() or f"exit code {completed.returncode}"
            return None, [f"PDF変換によるページ数再計算に失敗しました: {message}"]

        pdf_path = output_dir / f"{path.stem}.pdf"
        if not pdf_path.exists():
            pdf_candidates = list(output_dir.glob("*.pdf"))
            pdf_path = pdf_candidates[0] if pdf_candidates else pdf_path
        if not pdf_path.exists():
            return None, ["PDF変換後のPDFファイルが見つかりません"]

        try:
            reader = PdfReader(str(pdf_path))
            return len(reader.pages), []
        except Exception as exc:
            return None, [f"PDFページ数の読み取りに失敗しました: {exc}"]
